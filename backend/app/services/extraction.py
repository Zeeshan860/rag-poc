import io
import re
from dataclasses import dataclass
from pathlib import Path

import requests
from bs4 import BeautifulSoup
from docx import Document
from fastapi import HTTPException
from pypdf import PdfReader


@dataclass
class PageText:
    page_number: int | None
    text: str


@dataclass
class ExtractedDocument:
    pages: list[PageText]


def extract_pdf(data: bytes) -> list[PageText]:
    reader = PdfReader(io.BytesIO(data))
    return [
        PageText(page_number=page_num, text=page.extract_text() or "")
        for page_num, page in enumerate(reader.pages, start=1)
    ]


def extract_docx(data: bytes) -> str:
    document = Document(io.BytesIO(data))
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def extract_plain_text(data: bytes) -> str:
    return data.decode("utf-8")


def _extract_url_text(url: str) -> str:
    try:
        response = requests.get(
            url,
            timeout=15,
            headers={"User-Agent": "rag-poc/1.0"},
        )
        response.raise_for_status()
    except requests.RequestException as exc:
        raise HTTPException(status_code=502, detail=f"Failed to fetch URL: {exc}") from exc

    soup = BeautifulSoup(response.text, "html.parser")
    for tag in soup.find_all(["script", "style", "nav", "footer"]):
        tag.decompose()

    text = soup.get_text(separator="\n", strip=True)
    return re.sub(r"\n{3,}", "\n\n", text)


def extract_file_document(filename: str | None, data: bytes) -> ExtractedDocument:
    if not filename:
        raise HTTPException(status_code=400, detail="Filename is required")

    extension = Path(filename).suffix.lower()
    if extension == ".pdf":
        return ExtractedDocument(pages=extract_pdf(data))
    if extension == ".docx":
        return ExtractedDocument(pages=[PageText(page_number=None, text=extract_docx(data))])
    if extension in {".txt", ".md"}:
        return ExtractedDocument(pages=[PageText(page_number=None, text=extract_plain_text(data))])

    raise HTTPException(status_code=400, detail="Unsupported file type")


def extract_url_document(url: str) -> ExtractedDocument:
    return ExtractedDocument(pages=[PageText(page_number=None, text=_extract_url_text(url))])
